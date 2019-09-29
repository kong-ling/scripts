#!/usr/bin/python
# -*- coding: UTF-8 -*-
# coding=utf-8

import os
import sys
import re
import io
import docx

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
#sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='gb18030')
#os.system('chcp 936')

txt_files = (os.popen('dir /b *.docx')).readlines()

for f in txt_files:
    print(f, end='')

while True:
    matched_words = []
    #get user input
    print('\n============百宝箱===============')
    pattern = input('Input you option, Please: ')
    print('your input is "%s"\n' % pattern)

    for f in txt_files:
        f = f.strip()
        #print(f)
        #words = open(f, 'r', encoding='utf-8')
        words = docx.Document(f)
        print('paragraphs: ' + str(len(words.paragraphs)))
        for para in words.paragraphs:
            text = para.text
            #print(text)
            if re.search(pattern, text, re.IGNORECASE):
                matched_words.append(text)

    if len(matched_words):
        for idx in range(len(matched_words)):
            word = matched_words[idx].encode('utf-8')
            print("%3d:  %s" % (idx+1, word.decode('utf-8', 'ignore')))
    else:
        print('No word matches your input')
